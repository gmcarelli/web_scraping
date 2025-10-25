# Import necessary libraries
import undetected_chromedriver as uc
from bs4 import BeautifulSoup
import time
import os
import re
from docx import Document
from docx.shared import Pt

def get_essay_links(url):
    """
    Fetches the main page using undetected-chromedriver to bypass anti-bot measures.
    It extracts the theme of the essays and a list of links and titles for each essay.
    """
    links = []
    theme = "Default Theme" # Default theme name in case it's not found on the page
    try:
        options = uc.ChromeOptions()
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        options.add_argument('--disable-gpu')

        # Initialize the undetected Chrome driver
        driver = uc.Chrome(options=options)

        print("Accessing the main page...")
        driver.get(url)
        # Wait for the page to load completely, especially if it uses JavaScript
        time.sleep(5)

        # Parse the page content with BeautifulSoup
        soup = BeautifulSoup(driver.page_source, 'html.parser')

        # Close the browser
        driver.quit()

        # Extract the theme from the h2 title
        theme_element = soup.select_one("h2.container-title")
        if theme_element:
            theme = theme_element.get_text(strip=True)

        # Select all essay links based on their CSS selector
        essay_elements = soup.select("ul.itens-resp li a")
        for element in essay_elements:
            title = element.get_text(strip=True)
            link = element['href']
            if title and link:
                links.append({'title': title, 'url': link})

    except Exception as e:
        print(f"An error occurred while getting essay links: {e}")

    return theme, links

def get_essay_content(url):
    """
    Fetches an individual essay page and extracts the essay and correction text.
    It also removes any text that is styled in green.
    """
    essay_text = ""
    correction_text = ""
    try:
        options = uc.ChromeOptions()
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        options.add_argument('--disable-gpu')

        driver = uc.Chrome(options=options)
        driver.get(url)
        time.sleep(5)
        soup = BeautifulSoup(driver.page_source, 'html.parser')
        driver.quit()

        # Find and remove all spans with green color style, as requested
        for span in soup.find_all('span', style=lambda value: 'color: #008000' in value.lower() if value else False):
            span.decompose()

        # Select the main content body of the essay
        content = soup.select_one('.rt-body')
        if content:
            all_paragraphs = content.find_all('p')

            in_correction = False
            # Iterate through paragraphs to separate essay from correction
            for p in all_paragraphs:
                paragraph_text = p.get_text(strip=True)
                # The correction section starts with "comentários gerais"
                if "comentários gerais" in paragraph_text.lower():
                    in_correction = True

                if in_correction:
                    correction_text += paragraph_text + "\n\n"
                else:
                    essay_text += paragraph_text + "\n\n"

    except Exception as e:
        print(f"An error occurred while getting essay content from {url}: {e}")

    return essay_text.strip(), correction_text.strip()

def save_to_docx(theme, title, essay_text, correction_text):
    """
    Saves the extracted essay and correction to a .docx file inside a structured directory.
    """
    try:
        # Sanitize theme and title to make them valid directory and file names
        s_theme = re.sub(r'[\\/*?:"<>|]', "", theme)
        s_title = re.sub(r'[\\/*?:"<>|]', "", title)

        # Create the directory structure: essays -> theme_name
        dir_path = os.path.join("essays", s_theme)
        os.makedirs(dir_path, exist_ok=True)

        # Create a new Word document
        doc = Document()

        # Add the essay text to the first page
        doc.add_paragraph(essay_text)

        # Add a page break to separate essay and correction
        doc.add_page_break()

        # Add the correction text to the second page
        doc.add_paragraph(correction_text)

        # Define the full path for the output file and save it
        file_path = os.path.join(dir_path, f"{s_title}.docx")
        doc.save(file_path)
        print(f"  -> Successfully saved to {file_path}")

    except Exception as e:
        print(f"An error occurred while saving the .docx file for '{title}': {e}")

def main():
    """
    Main function to orchestrate the web scraping process.
    """
    # URL of the main page with the list of essays
    main_url = "https://educacao.uol.com.br/bancoderedacoes/propostas/qualificacao-e-o-futuro-do-emprego.htm"

    # Get the theme and all essay links from the main page
    theme, essay_links = get_essay_links(main_url)

    if essay_links:
        print(f"\nFound {len(essay_links)} essays for the theme: '{theme}'\n")
        # Process each essay link
        for i, essay in enumerate(essay_links, 1):
            print(f"Processing essay {i}/{len(essay_links)}: '{essay['title']}'...")
            # Get the content for the current essay
            essay_text, correction_text = get_essay_content(essay['url'])

            # If content was extracted successfully, save it to a .docx file
            if essay_text and correction_text:
                save_to_docx(theme, essay['title'], essay_text, correction_text)
            else:
                print("  -> Could not extract content, skipping file save.")
        print("\nScraping process completed!")
    else:
        print("No essays found. The script might be blocked by the website's security measures.")

# Entry point of the script
if __name__ == "__main__":
    main()
